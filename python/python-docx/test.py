from docx import Document

doc = Document('title.docx')

paragraphs = doc.paragraphs

tables = doc.tables

inline_shapes = doc.inline_shapes

for paragraph in paragraphs:
		print(paragraph.runs)
#     for run in paragraph.runs:
#         print(run.text)

# print(tables)
#
# print(inline_shapes)
#
# print(doc.part)
