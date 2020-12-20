import docx

def extract_docx_text(docFile):
    # 获取文档对象
    document = docx.Document(docFile)
    # 完整的text：
    docx_text = ""
    for para in document.paragraphs:
        docx_text += para.text + '\n'
    return docx_text


def extract_table_data(tables_i):
    columns = tables_i.columns
    """每列作为列表"""
    # tables_i_json = {}
    # # 列标题
    # columns_name_json = {}  # 用来下面表格的行做铺垫使用
    # for columns_i in range(len(columns)):
    #     tables_i_json[tables_i.cell(0, columns_i).text] = []
    #     columns_name_json[columns_i] = tables_i.cell(0, columns_i).text
    #
    # for row_i in range(1, len(tables_i.rows)):  # 从表格第二行开始循环读取表格数据
    #     for columns_i in range(len(columns)):
    #         tables_i_json[columns_name_json[columns_i]].append(tables_i.cell(row_i, columns_i).text)
    """每行一个json"""
    table_data = []
    # # 列标题
    head_name = []
    for columns_i in range(len(columns)):
        head_name.append(tables_i.cell(0, columns_i).text)

    for row_i in range(1, len(tables_i.rows)):  # 从表格第二行开始循环读取表格数据
        row_json = {}
        for columns_i in range(len(columns)):
            row_json[head_name[columns_i]] = tables_i.cell(row_i, columns_i).text
        table_data.append(row_json)

    return table_data



def extract_docx_excel_data(docFile):
    document = docx.Document(docFile)  # 读入文件
    # print("document",dir(document))
    tables = document.tables  # 获取文件中的表格集
    # print("dir", dir(document))
    # print("tables", tables)
    tables_json = []
    for tables_i in tables:
        tables_i_json = extract_table_data(tables_i)
        tables_json.append(tables_i_json)
    return tables_json

if __name__ == '__main__':
    docFile = 'newdocx.docx'
    # docx_text = extract_docx_text(docFile)
    # print("docx_text", docx_text)
    tables_json = extract_docx_excel_data(docFile)
    print(tables_json)